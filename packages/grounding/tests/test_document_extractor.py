"""DocumentExtractor: per-format text extraction for manual uploads (P6.1), best-effort."""

from io import BytesIO

from lunaris_grounding import DocumentExtractor


async def test_extracts_plain_text_file() -> None:
    # Act — a .txt (by content-type) is decoded; the title is the file stem.
    extracted = await DocumentExtractor().extract(
        filename="notes.txt", content_type="text/plain", data=b"Dijkstra relaxes edges."
    )

    # Assert
    assert extracted is not None
    assert extracted.text == "Dijkstra relaxes edges."
    assert extracted.title == "notes"


async def test_extracts_markdown_file() -> None:
    # Act — a .md (by extension, no content-type) is decoded.
    extracted = await DocumentExtractor().extract(
        filename="notes.md", content_type=None, data=b"# Heading\n\nSome **markdown** notes."
    )

    # Assert
    assert extracted is not None
    assert "markdown" in extracted.text
    assert extracted.title == "notes"


async def test_extracts_docx_paragraphs() -> None:
    # Arrange — build a real .docx in memory.
    from docx import Document

    buffer = BytesIO()
    document = Document()
    document.add_paragraph("Dijkstra relaxes edges.")
    document.add_paragraph("It finds shortest paths.")
    document.save(buffer)

    # Act
    extracted = await DocumentExtractor().extract(
        filename="paper.docx", content_type=None, data=buffer.getvalue()
    )

    # Assert — the paragraphs are concatenated; the title is the stem.
    assert extracted is not None
    assert "Dijkstra relaxes edges." in extracted.text
    assert "shortest paths" in extracted.text
    assert extracted.title == "paper"


async def test_unsupported_binary_type_returns_none() -> None:
    # Act — an unknown binary type is not guessed at.
    extracted = await DocumentExtractor().extract(
        filename="image.bin", content_type="application/octet-stream", data=b"\x00\x01\x02"
    )

    # Assert
    assert extracted is None


async def test_empty_text_returns_none() -> None:
    # Act — a blank text file yields nothing to ingest.
    extracted = await DocumentExtractor().extract(
        filename="blank.txt", content_type="text/plain", data=b"   \n  "
    )

    # Assert
    assert extracted is None


async def test_corrupt_pdf_degrades_to_none() -> None:
    # Act — bytes claiming to be a PDF but unparseable must not raise (best-effort).
    extracted = await DocumentExtractor().extract(
        filename="broken.pdf", content_type="application/pdf", data=b"not really a pdf"
    )

    # Assert
    assert extracted is None
