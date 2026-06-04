"""The default ``IDocumentExtractor``: per-format text extraction for manual uploads (P6.1).

Turns a user's uploaded file — ``.txt`` / ``.md`` (decoded directly), ``.pdf`` (pypdf), ``.docx``
(python-docx) — into plain text the corpus can chunk + embed. Parser deps are imported lazily so the
package still imports without them, and the sync parsers run off the event loop via
``asyncio.to_thread``. Extraction is best-effort: any failure (unsupported type, corrupt file, empty
result) collapses to ``None`` — the "skip this source" signal the caller uses.
"""

import asyncio
from io import BytesIO
from pathlib import Path

import structlog

from lunaris_grounding.ingest.extracted_document import ExtractedDocument

logger = structlog.get_logger()

_TEXT_EXTENSIONS = {".txt", ".md", ".markdown", ".text", ".rst"}


class DocumentExtractor:
    """Dispatches by file extension (then content-type) to a per-format text extractor."""

    async def extract(
        self, *, filename: str, content_type: str | None, data: bytes
    ) -> ExtractedDocument | None:
        suffix = Path(filename).suffix.lower()
        is_text = suffix in _TEXT_EXTENSIONS or (content_type or "").startswith("text/")
        if suffix == ".pdf" or content_type == "application/pdf":
            text = await asyncio.to_thread(_extract_pdf, data)
        elif suffix == ".docx":
            text = await asyncio.to_thread(_extract_docx, data)
        elif is_text:
            text = data.decode("utf-8", errors="replace")
        else:
            # Unknown binary types aren't guessed at — only the declared formats are extracted.
            logger.info("document_extract_unsupported", suffix=suffix, content_type=content_type)
            return None
        if not text or not text.strip():
            return None
        return ExtractedDocument(text=text, title=Path(filename).stem)


def _extract_pdf(data: bytes) -> str | None:
    """Concatenate the text of every PDF page (best-effort)."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(data))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:  # pypdf raises many undocumented types on corrupt files → skip
        logger.warning("document_extract_pdf_failed", error=type(exc).__name__)
        return None


def _extract_docx(data: bytes) -> str | None:
    """Concatenate the non-empty paragraphs of a .docx (best-effort)."""
    try:
        from docx import Document

        document = Document(BytesIO(data))
        return "\n\n".join(p.text for p in document.paragraphs if p.text.strip())
    except Exception as exc:  # python-docx raises many types on a corrupt/invalid file → skip
        logger.warning("document_extract_docx_failed", error=type(exc).__name__)
        return None
